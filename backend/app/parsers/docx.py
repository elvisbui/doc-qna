"""DOCX document parser using python-docx."""

from pathlib import Path

from docx import Document


def parse_docx(file_path: Path) -> str:
    """Extract text from a DOCX file including paragraphs and tables.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text content, separated by newlines.
    """
    doc = Document(str(file_path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)
